"""
Export API — download a generated assistant message as PDF or Word document.

Endpoints:
  GET /conversations/{conversation_id}/messages/{message_id}/export?format=pdf|docx
"""
import io
import logging
import re
from typing import Optional

from fastapi import APIRouter, Depends, Header, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from app.core.database import get_db
from app.core.security import decode_token
from app.models.models import Conversation, Message

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/conversations", tags=["export"])


def _user_from_token(authorization: Optional[str]) -> str:
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid token")
    payload = decode_token(authorization.removeprefix("Bearer "))
    if not payload or not payload.get("sub"):
        raise HTTPException(status_code=401, detail="Invalid token")
    return str(payload["sub"])


def _strip_markdown(text: str) -> str:
    """Very light markdown → plain text conversion for PDF/DOCX body."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)   # bold
    text = re.sub(r"\*(.*?)\*", r"\1", text)         # italic
    text = re.sub(r"`{1,3}(.*?)`{1,3}", r"\1", text, flags=re.S)  # code
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)  # headings
    return text


def _generate_pdf(title: str, subject: str, topic: str, content: str) -> bytes:
    """Generate a PDF file using ReportLab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=6,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1E293B"),
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#64748B"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=8,
        textColor=colors.HexColor("#1E293B"),
    )

    story = [
        Paragraph(title, title_style),
        Paragraph(f"{subject} · {topic}", subtitle_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E2E8F0")),
        Spacer(1, 0.2 * inch),
    ]

    for line in _strip_markdown(content).splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue
        # Encode & to avoid XML parsing errors in ReportLab
        line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(line, body_style))

    doc.build(story)
    return buf.getvalue()


def _generate_docx(title: str, subject: str, topic: str, content: str) -> bytes:
    """Generate a Word .docx file using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph(f"{subject} · {topic}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Body — detect headings by leading #
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            p = doc.add_paragraph(_strip_markdown(stripped))
            p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{conversation_id}/messages/{message_id}/export")
async def export_message(
    conversation_id: str,
    message_id: str,
    format: str = Query("pdf", regex="^(pdf|docx)$"),
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """
    Download an assistant message as **PDF** or **Word** document.

    Query param ``format``: ``pdf`` (default) or ``docx``.
    """
    user_id = _user_from_token(authorization)

    # Verify conversation ownership
    conv_res = await db.execute(
        select(Conversation).filter(
            Conversation.id == conversation_id,
            Conversation.user_id == user_id,
        )
    )
    conv = conv_res.scalar_one_or_none()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    # Fetch the message
    msg_res = await db.execute(
        select(Message).filter(
            Message.id == message_id,
            Message.conversation_id == conversation_id,
            Message.role == "assistant",
        )
    )
    msg = msg_res.scalar_one_or_none()
    if not msg:
        raise HTTPException(status_code=404, detail="Message not found")

    subject = conv.primary_subject or conv.subject or "General"
    topic = conv.primary_topic or conv.topic or "Untitled"
    content_type = msg.content_type or "content"
    title = f"{content_type.title()} — {topic}"

    try:
        if format == "pdf":
            data = _generate_pdf(title, subject, topic, msg.content or "")
            mime = "application/pdf"
            filename = f"{topic.replace(' ', '_')}_{content_type}.pdf"
        else:
            data = _generate_docx(title, subject, topic, msg.content or "")
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{topic.replace(' ', '_')}_{content_type}.docx"
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Export library not available: {exc}",
        )
    except Exception as exc:
        logger.error("Export error: %s", exc)
        raise HTTPException(status_code=500, detail="Export generation failed")

    return StreamingResponse(
        io.BytesIO(data),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
