"""
Document management and RAG pipeline endpoints
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Header, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
import uuid
import os
import tempfile
from typing import Optional
from app.core.database import get_db
from app.core.security import decode_token
from app.schemas import DocumentUpload, DocumentResponse, DocumentAnalysis
from app.models.models import Document, Chunk, User
from app.orchestration.content_orchestrator import ContentOrchestrator
from app.services.document_ingestion_service import DocumentIngestor
from app.core.config import settings
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])


def _extract_text_preview(file_path: str, file_ext: str, raw_content: bytes) -> str:
    """Extract a short text preview from the file for AI subject inference."""
    try:
        if file_ext == "pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc[:3]:  # First 3 pages
                text += page.get_text()
            doc.close()
            return text[:3000]
        elif file_ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs[:60])
            return text[:3000]
        else:  # txt and others
            return raw_content.decode("utf-8", errors="ignore")[:3000]
    except Exception as exc:
        logger.warning(f"Could not extract text preview ({exc}); falling back to raw decode")
        return raw_content.decode("utf-8", errors="ignore")[:3000]


def get_current_user_id(token: Optional[str]) -> str:
    """Extract user ID from token"""
    if not token or not token.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid token")
    
    token_str = token.replace("Bearer ", "")
    payload = decode_token(token_str)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    return payload.get("sub")


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    subject: str = None,
    topic: str = None,
    conversation_id: str = None,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a document for RAG ingestion
    
    Supports: PDF, DOCX, TXT
    """
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")
        
        user_id = get_current_user_id(authorization)
        
        # Validate file type
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in settings.ALLOWED_FILE_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed: {settings.ALLOWED_FILE_TYPES}"
            )
        
        # Validate file size
        content = await file.read()
        if len(content) > settings.MAX_DOCUMENT_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Max size: {settings.MAX_DOCUMENT_SIZE / 1024 / 1024:.0f}MB"
            )
        
        # Save file temporarily
        document_id = str(uuid.uuid4())
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # If no subject provided, try to infer it from document content using AI
            if not subject:
                try:
                    text_preview = _extract_text_preview(tmp_path, file_ext, content)
                    if text_preview.strip():
                        from app.services.llm_service import LLMService
                        subject = await LLMService.infer_subject(text_preview)
                        logger.info(f"Inferred subject for {file.filename!r}: {subject!r}")
                    else:
                        subject = "General"
                except Exception as infer_err:
                    logger.warning(f"Subject inference failed ({infer_err}); defaulting to 'General'")
                    subject = "General"

            # Extract full text from the document BEFORE ingestion (while tmp file exists)
            # This gives us clean, readable text for the RAG context instead of raw bytes.
            try:
                extracted_text = await DocumentIngestor.parse_file(tmp_path, file_ext)
                logger.info(
                    "Extracted %d chars of text from %s", len(extracted_text), file.filename
                )
            except Exception as extract_err:
                logger.warning(
                    "Full text extraction failed (%s); falling back to text preview", extract_err
                )
                extracted_text = _extract_text_preview(tmp_path, file_ext, content)

            # Ingest document through orchestrator
            ingestion_result = await ContentOrchestrator.ingest_teacher_document(
                document_id=document_id,
                file_path=tmp_path,
                file_type=file_ext,
                subject=subject or "General",
                topic=topic or "Unspecified",
                user_id=user_id,
                conversation_id=conversation_id,
            )

            # Save document metadata to DB — store extracted text, not raw bytes
            document = Document(
                id=document_id,
                user_id=user_id,
                conversation_id=conversation_id,
                filename=file.filename,
                file_type=file_ext,
                original_content=extracted_text[:50000],   # up to 50 k chars of clean text
                subject=subject,
                chunks_count=ingestion_result.get("chunks_created", 0),
                vector_index_ids=str(ingestion_result.get("chunk_ids", [])),
            )
            
            db.add(document)
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Document uploaded: {file.filename} ({document_id})")

            return document
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/fetch-url", response_model=DocumentResponse)
async def upload_from_url(
    request: dict,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """
    Fetch a document from a URL and ingest it for RAG.
    Request body: { url, subject, level, description? }
    """
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")

        user_id = get_current_user_id(authorization)
        url = request.get("url", "")
        subject = request.get("subject", "General")

        if not url:
            raise HTTPException(status_code=400, detail="URL is required")

        # Fetch content from URL
        content = await DocumentIngestor._parse_url(url)

        document_id = str(uuid.uuid4())

        # Ingest content
        ingestion_result = await ContentOrchestrator.ingest_teacher_document(
            document_id=document_id,
            file_path=url,
            file_type="url",
            subject=subject,
            topic=request.get("topic", "Unspecified"),
            user_id=user_id,
        )

        document = Document(
            id=document_id,
            user_id=user_id,
            filename=url[:200],
            file_type="url",
            original_content=content[:10000],
            chunks_count=ingestion_result.get("chunks_created", 0),
            vector_index_ids=str(ingestion_result.get("chunk_ids", [])),
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)

        logger.info(f"URL document ingested: {url} ({document_id})")
        return document

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching from URL: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ingest-text")
async def ingest_text_content(
    content: str,
    subject: str,
    topic: str = None,
    conversation_id: str = None,
    authorization: str = None,
    db: AsyncSession = Depends(get_db),
) -> DocumentAnalysis:
    """
    Ingest raw text content directly (without file upload)
    """
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")
        
        user_id = get_current_user_id(authorization)
        
        document_id = str(uuid.uuid4())
        
        # Ingest through orchestrator
        ingestion_result = await ContentOrchestrator.ingest_teacher_document(
            document_id=document_id,
            file_path="",
            file_type="txt",
            subject=subject,
            topic=topic or "Unspecified",
            user_id=user_id,
            conversation_id=conversation_id,
        )
        
        # Save to DB
        document = Document(
            id=document_id,
            user_id=user_id,
            conversation_id=conversation_id,
            filename=f"text_input_{document_id[:8]}.txt",
            file_type="txt",
            original_content=content[:10000],
            chunks_count=ingestion_result.get("chunks_created", 0),
            vector_index_ids=str(ingestion_result.get("chunk_ids", [])),
        )
        
        db.add(document)
        await db.commit()
        
        logger.info(f"Text content ingested: {document_id}")
        
        return DocumentAnalysis(
            document_id=document_id,
            chunks_created=ingestion_result.get("chunks_created", 0),
            chunk_ids=ingestion_result.get("chunk_ids", []),
            subject=subject,
            topic=topic,
            status="success",
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error ingesting text: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("", response_model=list[DocumentResponse])
async def list_documents(
    conversation_id: str = None,
    authorization: str = None,
    db: AsyncSession = Depends(get_db),
):
    """List documents uploaded by user"""
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")
        
        user_id = get_current_user_id(authorization)
        
        query = select(Document).filter(Document.user_id == user_id)
        
        if conversation_id:
            query = query.filter(Document.conversation_id == conversation_id)
        
        result = await db.execute(query)
        documents = result.scalars().all()
        
        return documents
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{document_id}", status_code=204)
async def delete_document(
    document_id: str,
    authorization: str = None,
    db: AsyncSession = Depends(get_db),
):
    """Delete a document"""
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")
        
        user_id = get_current_user_id(authorization)
        
        # Verify ownership
        result = await db.execute(
            select(Document).filter(
                Document.id == document_id,
                Document.user_id == user_id
            )
        )
        document = result.scalar_one_or_none()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        await db.delete(document)
        await db.commit()
        
        logger.info(f"Document deleted: {document_id}")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """Get a specific document by ID"""
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")

        user_id = get_current_user_id(authorization)

        result = await db.execute(
            select(Document).filter(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        return document

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{document_id}/chunks")
async def get_document_chunks(
    document_id: str,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """
    Retrieve the individual text chunks for a document.

    Useful for debugging the RAG pipeline: inspect how the document was
    chunked and what text each chunk contains.
    """
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")

        user_id = get_current_user_id(authorization)

        # Verify ownership
        doc_result = await db.execute(
            select(Document).filter(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )
        document = doc_result.scalar_one_or_none()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Fetch stored chunks
        chunk_result = await db.execute(
            select(Chunk)
            .filter(Chunk.document_id == document_id)
            .order_by(Chunk.chunk_index)
        )
        chunks = chunk_result.scalars().all()

        return {
            "document_id": document_id,
            "filename": document.filename,
            "total_chunks": len(chunks),
            "chunks": [
                {
                    "id": c.id,
                    "chunk_index": c.chunk_index,
                    "text": c.text,
                    "chunk_size": c.chunk_size,
                    "overlap_info": c.overlap_info,
                    "vector_id": c.vector_id,
                }
                for c in chunks
            ],
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching chunks for document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/syllabus", response_model=DocumentAnalysis)
async def ingest_syllabus(
    request: dict,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """
    Accept a structured course syllabus and ingest it as a RAG document.

    Request body::

        {
          "course_name": "Introduction to Biology",
          "subject": "Biology",
          "week": 3,
          "learning_objectives": ["Understand cell structure", "Apply membrane concepts"],
          "topics": [
            {"name": "Cell membrane", "bloom_level": "Understand"},
            {"name": "Osmosis and diffusion", "bloom_level": "Apply"}
          ],
          "notes": "Focus on transport mechanisms for the exam"
        }

    The backend formats the fields into a human-readable text document and
    ingests it via the standard RAG pipeline so it becomes retrievable context
    for future content generation requests.
    """
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing authorization")
    user_id = get_current_user_id(authorization)

    course_name = request.get("course_name", "Untitled Course")
    subject = request.get("subject", "General")
    week = request.get("week")
    objectives = request.get("learning_objectives", [])
    topics = request.get("topics", [])
    notes = request.get("notes", "")

    # Build a structured plain-text document
    lines = [
        f"# Course Syllabus: {course_name}",
        f"Subject: {subject}",
    ]
    if week:
        lines.append(f"Week: {week}")

    if objectives:
        lines.append("\n## Learning Objectives")
        for obj in objectives:
            lines.append(f"- {obj}")

    if topics:
        lines.append("\n## Topics Covered")
        for t in topics:
            if isinstance(t, dict):
                bloom = t.get("bloom_level", "")
                name = t.get("name", str(t))
                lines.append(f"- {name}" + (f"  [{bloom}]" if bloom else ""))
            else:
                lines.append(f"- {t}")

    if notes:
        lines.append(f"\n## Notes\n{notes}")

    content = "\n".join(lines)

    document_id = str(uuid.uuid4())
    filename = f"syllabus_{course_name[:40].replace(' ', '_')}_{document_id[:8]}.txt"

    try:
        ingestion_result = await ContentOrchestrator.ingest_teacher_document(
            document_id=document_id,
            file_path="",
            file_type="txt",
            subject=subject,
            topic=course_name,
            user_id=user_id,
        )
    except Exception as e:
        logger.warning("Syllabus vector ingestion failed: %s", e)
        ingestion_result = {"chunks_created": 0, "chunk_ids": []}

    document = Document(
        id=document_id,
        user_id=user_id,
        filename=filename,
        file_type="txt",
        original_content=content,
        subject=subject,
        chunks_count=ingestion_result.get("chunks_created", 0),
        vector_index_ids=str(ingestion_result.get("chunk_ids", [])),
    )
    db.add(document)
    await db.commit()

    logger.info("Syllabus document ingested: %s for user %s", document_id, user_id)

    return DocumentAnalysis(
        document_id=document_id,
        chunks_created=ingestion_result.get("chunks_created", 0),
        chunk_ids=ingestion_result.get("chunk_ids", []),
        subject=subject,
        topic=course_name,
        status="success",
    )


@router.post("/search/semantic")
async def semantic_search(
    request: dict,
    authorization: Optional[str] = Header(default=None),
    db: AsyncSession = Depends(get_db),
):
    """Perform semantic search over indexed documents"""
    try:
        if not authorization:
            raise HTTPException(status_code=401, detail="Missing authorization")

        from app.services.vector_service import VectorDatabaseService

        query = request.get("query", "")
        subject = request.get("subject")
        topic = request.get("topic")
        limit = request.get("limit", 5)

        results = await VectorDatabaseService.query(
            query_text=query,
            n_results=limit,
            where_filter={"subject": subject} if subject else None,
        )

        chunks = []
        for idx, (doc, meta) in enumerate(
            zip(
                results.get("results", []),
                results.get("metadatas", [{}] * limit),
            )
        ):
            chunks.append({
                "chunkIndex": idx,
                "text": doc.get("content", ""),
                "chunkSize": len(doc.get("content", "")),
                "similarityScore": doc.get("relevance_score", 0.0),
                "metadata": meta,
            })

        return {
            "chunks": chunks,
            "totalFound": len(chunks),
            "query": query,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error during semantic search: {e}")
        raise HTTPException(status_code=500, detail=str(e))
